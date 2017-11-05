import sys
import urllib2
import copy
import collections
from district import *
from docx import Document
from docx.shared import Inches
from HTMLParser import HTMLParser

mindcp = int(sys.argv[1])
h = HTMLParser()
document = Document()
d36 = District(36)
document.add_heading('D36 DCP Points Leaders - %s' % (d36.get_dashdate()), 0)
d36dcparea = d36.dcpsort()
for ndcp in range(10,mindcp-1,-1):
   if len(d36dcparea[ndcp]) > 0:
      document.add_heading('%d DCP Points' % (ndcp), level=1)
      table = document.add_table(rows=1, cols=3)
      hdr_cells1 = table.rows[0].cells
      hdr_cells1[0].text = 'Area'
      hdr_cells1[1].text = 'Club Number'
      hdr_cells1[2].text = 'Name'
      sys.stdout.write('\n%d DCP Points:\n' % (ndcp))
      sys.stdout.write('-------------\n')
      for c in d36dcparea[ndcp]:
         row_cells = table.add_row().cells
         row_cells[0].text = str(c.area)
         row_cells[1].text = str(c.number)
         row_cells[2].text = h.unescape(c.name)
         sys.stdout.write('%d   % 8d      %s\n' % (c.area, c.number, c.name))
outfname = 'dcppoints_%s.docx' % (d36.get_dashdate())
print 'output file name = ', outfname
document.save(outfname)
