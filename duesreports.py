import sys
import urllib2
import copy
import collections
from district import *
import PyPDF2
from docx import Document
from docx.shared import Inches
from HTMLParser import HTMLParser

h = HTMLParser()
d36 = District(36)
print 'd36 = ', d36.__dict__
dl = d36.divs.keys()
dl.sort()
document = Document()
document.add_heading('D36 Late Dues Report - %s' % (d36.get_dashdate()), 0)
for d in dl:
   document.add_heading('Division %s' % (d), level=1)
   print '\nDivision %s' % (d)
   al = d36.divs[d].keys()
   al.sort()
   firstarea = True
   for a in al:
      nlate = 0
      for c in d36.divs[d][a]:
         if c.october < 8:
            nlate += 1
      if nlate > 0:
         print '   Area ', a
         document.add_heading('Area %s' % (a), 2)
         table = document.add_table(rows=1, cols=3)
         hdr_cells1 = table.rows[0].cells
         hdr_cells1[0].text = 'Club'
         hdr_cells1[1].text = 'Base'
         hdr_cells1[2].text = 'October'
         for c in d36.divs[d][a]:
            if c.october < 8:
               row_cells = table.add_row().cells
               row_cells[0].text = h.unescape(c.name)
               row_cells[1].text = str(c.nbase)
               row_cells[2].text = str(c.october)
               firstarea = True
               print '      %-58s   %2d   %2d' % (c.name, c.nbase, c.october)
   if d != dl[-1]:
      document.add_page_break()

outfname = 'latedues_%s.docx' % (d36.get_dashdate())
print 'output file name = ', outfname
document.save(outfname)
